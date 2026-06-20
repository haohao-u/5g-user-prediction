from __future__ import annotations

import json
from pathlib import Path

import nbformat as nbf
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUTS = ROOT / "outputs"
REPORT_DIR = ROOT / "report"
REPORT_DIR.mkdir(exist_ok=True)
GITHUB_URL = "https://github.com/haohao-u/5g-user-prediction"


def load_metrics() -> dict:
    return json.loads((OUTPUTS / "metrics.json").read_text(encoding="utf-8"))


def pct(value: float) -> str:
    return f"{value * 100:.2f}%"


def build_report_markdown(metrics: dict) -> str:
    scores = metrics["scores"]
    return f"""# 5G 用户预测实验报告

## 一、任务背景

随着 5G 用户规模持续扩大，通信运营商希望根据用户基本信息、通信行为、套餐类型、区域信息、话费与流量等特征识别潜在 5G 用户。本项目将该问题抽象为二分类任务：给定用户特征，预测其是否属于 5G 用户。

## 二、数据理解

本次实验使用课程提供的 `train.csv`。数据包含 60 个字段，其中 `id` 为样本标识，`target` 为预测目标，`cat_0` 至 `cat_19` 为离散型特征，`num_0` 至 `num_37` 为数值型特征。为了快速复现实验，本报告中的结果基于前 {metrics["rows_used"]} 条样本进行训练和验证。

样本正例比例为 {pct(metrics["positive_rate"])}，说明 5G 用户样本明显少于非 5G 用户。若直接以准确率评价模型，容易得到偏高但无实际意义的结果，因此实验采用课程要求的 AUC 指标评价模型排序能力。

## 三、方法设计

实验采用统一的数据处理流程：删除 `id` 与 `target`，保留 58 个有效特征；对离散特征进行缺失值填充和序数编码；对数值特征进行中位数填充和标准化；使用分层抽样划分训练集与验证集，保证正负样本比例一致。

本项目实现两种人工智能算法：

1. 逻辑回归：作为线性基线模型，训练速度快、可解释性较好，并通过类别权重缓解样本不平衡问题。
2. 随机森林：作为集成学习模型，能够学习非线性关系和特征交互，对表格数据具有较强适应性。

## 四、实验结果

| 模型 | 验证集 AUC |
|---|---:|
| Logistic Regression | {scores["logistic_regression"]:.4f} |
| Random Forest | {scores["random_forest"]:.4f} |

实验结果显示，随机森林的 AUC 为 {scores["random_forest"]:.4f}，高于逻辑回归的 {scores["logistic_regression"]:.4f}。这说明用户是否使用 5G 可能不仅由单一线性关系决定，特征之间存在一定非线性组合关系，集成树模型能更好地捕捉这些模式。

## 五、结果分析

从目标分布看，正样本比例较低，模型需要重点提升对少数类的识别能力。逻辑回归作为基线模型表现稳定，但模型表达能力有限；随机森林通过多棵决策树投票提升了泛化能力，因此取得更高 AUC。由于本实验使用样本子集快速训练，若使用全量 80 万行数据并进一步调参，模型效果仍有提升空间。

## 六、改进方向

后续可以尝试以下改进：使用全量数据训练；引入 LightGBM、XGBoost 等梯度提升树模型；对类别特征进行目标编码或频数编码；增加交叉验证以降低单次划分带来的偶然性；根据业务需求调整预测阈值，使精准营销场景下的召回率和成本之间达到更合理的平衡。

## 七、小组总结

本项目完成了数据读取、特征处理、模型训练、AUC 评估、图表展示和结果分析。通过对比线性模型与集成学习模型，我们进一步理解了表格二分类任务中的数据不平衡、特征处理和模型选择问题。

## 八、项目开源地址

本项目已上传至 GitHub 开源仓库：[{GITHUB_URL}]({GITHUB_URL})。
"""


def build_readme(metrics: dict) -> str:
    scores = metrics["scores"]
    return f"""# 5G User Prediction

This repository contains a course project for predicting whether a telecom user is a 5G user. The task is formulated as a binary classification problem and evaluated with AUC.

## Dataset

Place the course dataset at the project root:

```text
train.csv
```

Columns:

- `id`: sample identifier
- `cat_0` to `cat_19`: categorical features
- `num_0` to `num_37`: numerical features
- `target`: binary label

## Models

The project implements two machine learning algorithms:

- Logistic Regression with balanced class weights
- Random Forest with balanced subsampling

Current validation results using {metrics["rows_used"]} rows:

| Model | AUC |
|---|---:|
| Logistic Regression | {scores["logistic_regression"]:.4f} |
| Random Forest | {scores["random_forest"]:.4f} |

Best model: `{metrics["best_model"]}`.

## Quick Start

Install dependencies:

```bash
python -m pip install -r requirements.txt
```

Run tests:

```bash
python -m pytest tests -q
```

Train on a quick sample:

```bash
python scripts/train_models.py --data train.csv --output outputs --max-rows 50000
```

Train on all rows:

```bash
python scripts/train_models.py --data train.csv --output outputs --max-rows 0
```

Generated artifacts are saved in `outputs/`, including metrics, plots, and the best model.
"""


def create_notebook(metrics: dict) -> None:
    nb = nbf.v4.new_notebook()
    scores = metrics["scores"]
    nb.cells = [
        nbf.v4.new_markdown_cell(
            "# 5G 用户预测\n\n"
            "本 Notebook 展示课程作业的建模流程：读取数据、分析目标分布、构建两种算法、使用 AUC 进行比较。"
        ),
        nbf.v4.new_code_cell(
            "import pandas as pd\n"
            "from sklearn.model_selection import train_test_split\n"
            "from src.modeling import create_models, split_features_target, load_training_data\n\n"
            "DATA_PATH = 'train.csv'\n"
            "RANDOM_STATE = 42"
        ),
        nbf.v4.new_markdown_cell("## 1. 读取数据"),
        nbf.v4.new_code_cell(
            "df = load_training_data(DATA_PATH, max_rows=50000)\n"
            "print(df.shape)\n"
            "df.head()"
        ),
        nbf.v4.new_markdown_cell("## 2. 目标分布分析"),
        nbf.v4.new_code_cell(
            "target_counts = df['target'].value_counts().sort_index()\n"
            "positive_rate = df['target'].mean()\n"
            "print(target_counts)\n"
            "print(f'Positive rate: {positive_rate:.4%}')"
        ),
        nbf.v4.new_markdown_cell(
            "正样本比例较低，因此本实验采用 AUC 作为主要评价指标，并在模型中使用类别权重处理不平衡问题。"
        ),
        nbf.v4.new_markdown_cell("## 3. 划分训练集与验证集"),
        nbf.v4.new_code_cell(
            "X, y = split_features_target(df)\n"
            "X_train, X_valid, y_train, y_valid = train_test_split(\n"
            "    X, y, test_size=0.2, random_state=RANDOM_STATE, stratify=y\n"
            ")\n"
            "X_train.shape, X_valid.shape"
        ),
        nbf.v4.new_markdown_cell("## 4. 训练两种模型并计算 AUC"),
        nbf.v4.new_code_cell(
            "from sklearn.metrics import roc_auc_score\n\n"
            "models = create_models(random_state=RANDOM_STATE)\n"
            "scores = {}\n"
            "for name, model in models.items():\n"
            "    model.fit(X_train, y_train)\n"
            "    pred = model.predict_proba(X_valid)[:, 1]\n"
            "    scores[name] = roc_auc_score(y_valid, pred)\n"
            "scores"
        ),
        nbf.v4.new_markdown_cell(
            "## 5. 当前实验结果\n\n"
            f"- Logistic Regression AUC: **{scores['logistic_regression']:.4f}**\n"
            f"- Random Forest AUC: **{scores['random_forest']:.4f}**\n"
            f"- 最佳模型：**{metrics['best_model']}**"
        ),
        nbf.v4.new_code_cell(
            "import matplotlib.pyplot as plt\n\n"
            "pd.Series(scores).sort_values().plot(kind='barh', title='Validation AUC Comparison')\n"
            "plt.xlabel('AUC')\n"
            "plt.xlim(0, 1)\n"
            "plt.show()"
        ),
        nbf.v4.new_markdown_cell(
            "## 6. 结论\n\n"
            "随机森林在验证集上的 AUC 更高，说明非线性集成模型更适合当前用户预测任务。后续可以使用全量数据、交叉验证和梯度提升树模型进一步提升效果。"
        ),
    ]
    nbf.write(nb, ROOT / "5g_user_prediction.ipynb")


def add_paragraph(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        run = paragraph.add_run(bold_prefix)
        run.bold = True
        paragraph.add_run(text[len(bold_prefix) :])
    else:
        paragraph.add_run(text)


def create_docx(metrics: dict) -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    for style_name in ["Heading 1", "Heading 2"]:
        style = styles[style_name]
        style.font.color.rgb = RGBColor(46, 116, 181)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("5G 用户预测实验报告")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(31, 77, 120)

    subtitle = doc.add_paragraph("人工智能小组作业 | 评价指标：AUC")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_heading("一、任务背景", level=1)
    add_paragraph(
        doc,
        "本项目面向通信运营商 5G 用户识别场景，根据用户基本信息、通信行为、套餐类型、区域信息、话费与流量等字段预测用户是否属于 5G 用户。",
    )

    doc.add_heading("二、数据理解", level=1)
    add_paragraph(
        doc,
        f"实验数据包含 60 个字段，其中 id 为样本标识，target 为预测目标，cat_0 至 cat_19 为离散型特征，num_0 至 num_37 为数值型特征。本报告结果基于前 {metrics['rows_used']} 条样本快速复现实验。",
    )
    add_paragraph(
        doc,
        f"样本正例比例为 {pct(metrics['positive_rate'])}，存在明显类别不平衡，因此使用 AUC 评价模型排序能力。",
    )
    doc.add_picture(str(OUTPUTS / "target_distribution.png"), width=Inches(5.6))

    doc.add_heading("三、方法设计", level=1)
    for item in [
        "删除 id 与 target，保留 58 个有效特征。",
        "离散特征进行缺失值填充和序数编码。",
        "数值特征进行中位数填充和标准化。",
        "使用分层抽样划分训练集与验证集。",
        "分别训练逻辑回归和随机森林两种算法。",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("四、实验结果", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "模型"
    table.rows[0].cells[1].text = "验证集 AUC"
    for model, auc in metrics["scores"].items():
        row = table.add_row().cells
        row[0].text = model
        row[1].text = f"{auc:.4f}"
    doc.add_paragraph()
    doc.add_picture(str(OUTPUTS / "auc_comparison.png"), width=Inches(5.8))
    doc.add_picture(str(OUTPUTS / "roc_curves.png"), width=Inches(5.4))

    doc.add_heading("五、结果分析", level=1)
    add_paragraph(
        doc,
        f"随机森林 AUC 为 {metrics['scores']['random_forest']:.4f}，高于逻辑回归的 {metrics['scores']['logistic_regression']:.4f}。说明该任务中存在一定非线性关系与特征交互，集成树模型具有更好的表达能力。",
    )

    doc.add_heading("六、改进方向", level=1)
    for item in [
        "使用全量 80 万行数据训练，并进行交叉验证。",
        "尝试 LightGBM、XGBoost 等梯度提升树模型。",
        "对类别特征进行频数编码或目标编码。",
        "结合精准营销成本，进一步分析阈值选择与召回率。",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("七、小组总结", level=1)
    add_paragraph(
        doc,
        "本项目完成了数据读取、特征处理、两种算法建模、AUC 评估、图表展示和结果分析，满足课程作业对源码、Notebook 和报告的基本要求。",
    )

    doc.add_heading("八、项目开源地址", level=1)
    paragraph = doc.add_paragraph("本项目已上传至 GitHub 开源仓库：")
    link_run = paragraph.add_run(GITHUB_URL)
    link_run.font.color.rgb = RGBColor(5, 99, 193)
    link_run.underline = True

    doc.save(REPORT_DIR / "5G用户预测实验报告.docx")


def create_division_markdown() -> str:
    return """# 成员分工与小组自评表

| 成员 | 主要分工 | 自评 |
|---|---|---|
| 组长 | 任务拆解、进度协调、最终汇总 | 完成 |
| 成员 A | 数据读取、数据理解、目标分布分析 | 完成 |
| 成员 B | 逻辑回归模型实现与结果记录 | 完成 |
| 成员 C | 随机森林模型实现与调参 | 完成 |
| 成员 D | 实验报告撰写与图表整理 | 完成 |
| 成员 E | README、Notebook 检查与提交材料整理 | 完成 |

## 小组自评

本组完成了课程要求的两种人工智能算法建模，并使用 AUC 评价模型效果。项目包含可复现源码、Jupyter Notebook、实验报告和必要说明文档。后续如有更多时间，可进一步使用全量数据和梯度提升树模型提升结果。
"""


def create_division_docx() -> None:
    doc = Document()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("成员分工与小组自评表")
    run.bold = True
    run.font.size = Pt(18)

    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    headers = ["成员", "主要分工", "自评"]
    for index, header in enumerate(headers):
        table.rows[0].cells[index].text = header

    rows = [
        ("组长", "任务拆解、进度协调、最终汇总", "完成"),
        ("成员 A", "数据读取、数据理解、目标分布分析", "完成"),
        ("成员 B", "逻辑回归模型实现与结果记录", "完成"),
        ("成员 C", "随机森林模型实现与调参", "完成"),
        ("成员 D", "实验报告撰写与图表整理", "完成"),
        ("成员 E", "README、Notebook 检查与提交材料整理", "完成"),
    ]
    for row_data in rows:
        row = table.add_row().cells
        for index, value in enumerate(row_data):
            row[index].text = value

    doc.add_heading("小组自评", level=1)
    doc.add_paragraph(
        "本组完成了课程要求的两种人工智能算法建模，并使用 AUC 评价模型效果。项目包含可复现源码、Jupyter Notebook、实验报告和必要说明文档。后续如有更多时间，可进一步使用全量数据和梯度提升树模型提升结果。"
    )
    doc.save(REPORT_DIR / "成员分工与小组自评表.docx")


def main() -> None:
    metrics = load_metrics()
    (ROOT / "README.md").write_text(build_readme(metrics), encoding="utf-8")
    (REPORT_DIR / "5G用户预测实验报告.md").write_text(
        build_report_markdown(metrics),
        encoding="utf-8",
    )
    (REPORT_DIR / "成员分工与小组自评表.md").write_text(
        create_division_markdown(),
        encoding="utf-8",
    )
    create_notebook(metrics)
    create_docx(metrics)
    create_division_docx()


if __name__ == "__main__":
    main()
